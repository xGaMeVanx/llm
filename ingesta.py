"""
Ingesta de la biblioteca de LLMario — genera indice.sqlite.

Lee archivos de conocimiento desde biblioteca/ y docs/ en estos formatos:
PDF, EPUB, DOCX, DOC, TXT y Markdown. Extrae el texto, lo trocea en
fragmentos con solapamiento y lo embebe con Gemini (gemini-embedding-001).
El resultado es indice.sqlite, una base SQLite con búsqueda vectorial
(sqlite-vec) que main.py consulta sin cargar todo en RAM.

Esquema:
  fragmentos      (id, fuente, titulo, pagina, texto)
  vec_fragmentos  tabla vec0: embedding float[768], coseno

Nota sobre .doc: el formato binario antiguo de Word no se lee directamente;
se convierte con LibreOffice (soffice) si está instalado. Si no, convierte el
archivo a .docx o .pdf y vuelve a correr la ingesta.

Uso:
    export GEMINI_API_KEY=...
    python ingesta.py              # genera indice.sqlite
    python ingesta.py --solo-probar  # muestra el troceado sin gastar la API
"""

import argparse
import os
import pathlib
import shutil
import sqlite3
import subprocess
import sys
import time

import numpy as np
import sqlite_vec
from google import genai
from google.genai import types

DIRECTORIO = pathlib.Path(__file__).parent
CARPETAS = [DIRECTORIO / "biblioteca", DIRECTORIO / "docs"]
EXTENSIONES = {".pdf", ".epub", ".docx", ".doc", ".txt", ".md"}
ARCHIVO_SQLITE = DIRECTORIO / "indice.sqlite"
MODELO = "gemini-embedding-001"
DIMENSIONES = 768
MAX_CHAR_FRAGMENTO = 1200
SOLAPE = 200  # caracteres que se repiten entre fragmentos vecinos


# ---------------------------------------------------------------------------
# Extracción de texto por formato
# ---------------------------------------------------------------------------

def _leer_pdf(ruta):
    from pypdf import PdfReader

    bloques = []
    lector = PdfReader(str(ruta))
    for numero, pagina in enumerate(lector.pages, 1):
        texto = (pagina.extract_text() or "").strip()
        if texto:
            bloques.append({"titulo": f"página {numero}", "pagina": numero, "texto": texto})
    return bloques


def _leer_epub(ruta):
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup

    bloques = []
    libro = epub.read_epub(str(ruta))
    for item in libro.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        html = item.get_content().decode("utf-8", errors="ignore")
        sopa = BeautifulSoup(html, "lxml")
        encabezado = sopa.find(["h1", "h2"])
        titulo = encabezado.get_text(strip=True) if encabezado else "capítulo"
        lineas = [l.strip() for l in sopa.get_text("\n").splitlines() if l.strip()]
        if lineas:
            bloques.append({"titulo": titulo, "pagina": None, "texto": "\n".join(lineas)})
    return bloques


def _leer_docx(ruta):
    import docx

    documento = docx.Document(str(ruta))
    parrafos = [p.text.strip() for p in documento.paragraphs if p.text.strip()]
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if celda.text.strip():
                    parrafos.append(celda.text.strip())
    texto = "\n\n".join(parrafos)
    return [{"titulo": "documento", "pagina": None, "texto": texto}] if texto else []


def _convertir_doc_a_docx(ruta):
    """Convierte .doc a .docx con LibreOffice. Devuelve la ruta o None."""
    soffice = shutil.which("soffice")
    if not soffice:
        return None
    destino = ruta.with_suffix(".docx")
    subprocess.run(
        [
            soffice, "--headless", "--convert-to", "docx",
            "--outdir", str(ruta.parent), str(ruta),
        ],
        capture_output=True,
        timeout=120,
    )
    return destino if destino.exists() else None


def _leer_doc(ruta):
    convertido = _convertir_doc_a_docx(ruta)
    if convertido is None:
        print(
            f"AVISO: no pude leer {ruta.name} (.doc). "
            "Instala LibreOffice o conviértelo a .docx/.pdf."
        )
        return []
    bloques = _leer_docx(convertido)
    convertido.unlink(missing_ok=True)
    return bloques


def _leer_texto(ruta):
    for codificacion in ("utf-8", "latin-1"):
        try:
            return ruta.read_text(encoding=codificacion)
        except UnicodeDecodeError:
            continue
    return ""


def _bloques_markdown(ruta):
    """Secciones '## ' de un .md."""
    lineas = ruta.read_text(encoding="utf-8").splitlines()
    secciones = []
    titulo_actual = None
    cuerpo = []
    for linea in lineas:
        if linea.startswith("## "):
            if titulo_actual is not None:
                secciones.append((titulo_actual, cuerpo))
            titulo_actual = linea[3:].strip()
            cuerpo = []
        elif titulo_actual is not None:
            cuerpo.append(linea)
    if titulo_actual is not None:
        secciones.append((titulo_actual, cuerpo))
    return [
        {"titulo": titulo, "pagina": None, "texto": "\n".join(cuerpo).strip()}
        for titulo, cuerpo in secciones
        if "\n".join(cuerpo).strip()
    ]


def extraer_archivo(ruta):
    """Devuelve los bloques [{titulo, pagina, texto}] de un archivo."""
    extension = ruta.suffix.lower()
    if extension == ".pdf":
        return _leer_pdf(ruta)
    if extension == ".epub":
        return _leer_epub(ruta)
    if extension == ".docx":
        return _leer_docx(ruta)
    if extension == ".doc":
        return _leer_doc(ruta)
    if extension == ".md":
        return _bloques_markdown(ruta)
    if extension == ".txt":
        texto = _leer_texto(ruta)
        return [{"titulo": "documento", "pagina": None, "texto": texto}] if texto.strip() else []
    return []


# ---------------------------------------------------------------------------
# Troceado con solapamiento
# ---------------------------------------------------------------------------

def _trocear(texto):
    """Parte un texto en fragmentos de hasta MAX_CHAR_FRAGMENTO con solape."""
    if len(texto) <= MAX_CHAR_FRAGMENTO:
        return [texto]

    parrafos = [p.strip() for p in texto.split("\n\n") if p.strip()]

    # Los párrafos gigantes se parten con ventana deslizante + solape.
    planos = []
    for parrafo in parrafos:
        while len(parrafo) > MAX_CHAR_FRAGMENTO:
            planos.append(parrafo[:MAX_CHAR_FRAGMENTO])
            parrafo = parrafo[MAX_CHAR_FRAGMENTO - SOLAPE:]
        if parrafo:
            planos.append(parrafo)

    piezas, actual = [], ""
    for parrafo in planos:
        if actual and len(actual) + len(parrafo) + 2 > MAX_CHAR_FRAGMENTO:
            piezas.append(actual)
            actual = ""
        actual = f"{actual}\n\n{parrafo}" if actual else parrafo
    if actual:
        piezas.append(actual)
    return piezas


def _trocear_bloques(ruta, bloques):
    fragmentos = []
    for bloque in bloques:
        for trozo in _trocear(bloque["texto"]):
            fragmentos.append(
                {
                    "fuente": ruta.name,
                    "titulo": bloque["titulo"],
                    "pagina": bloque["pagina"],
                    "texto": trozo,
                }
            )
    return fragmentos


def _archivos():
    for carpeta in CARPETAS:
        if not carpeta.exists():
            continue
        for ruta in sorted(carpeta.rglob("*")):
            if ruta.name.startswith("~$") or ruta.name.startswith("."):
                continue
            if ruta.suffix.lower() in EXTENSIONES:
                yield ruta


# ---------------------------------------------------------------------------
# Embedding y escritura a SQLite
# ---------------------------------------------------------------------------

def embeber(cliente, texto):
    """Embedding normalizado de un texto con Gemini."""
    respuesta = cliente.models.embed_content(
        model=MODELO,
        contents=texto,
        config=types.EmbedContentConfig(
            task_type="RETRIEVAL_DOCUMENT",
            output_dimensionality=DIMENSIONES,
        ),
    )
    vector = np.array(respuesta.embeddings[0].values, dtype=np.float32)
    return vector / np.linalg.norm(vector)


def _abrir_db():
    """Crea indice.sqlite desde cero con el esquema de fragmentos + vectores."""
    ARCHIVO_SQLITE.unlink(missing_ok=True)
    db = sqlite3.connect(ARCHIVO_SQLITE)
    db.enable_load_extension(True)
    sqlite_vec.load(db)
    db.enable_load_extension(False)
    db.execute(
        "CREATE TABLE fragmentos ("
        " id INTEGER PRIMARY KEY,"
        " fuente TEXT NOT NULL,"
        " titulo TEXT,"
        " pagina INTEGER,"
        " texto TEXT NOT NULL)"
    )
    db.execute(
        f"CREATE VIRTUAL TABLE vec_fragmentos USING "
        f"vec0(embedding float[{DIMENSIONES}] distance_metric=cosine)"
    )
    return db


def _insertar(db, fragmento, vector):
    cursor = db.execute(
        "INSERT INTO fragmentos (fuente, titulo, pagina, texto) VALUES (?, ?, ?, ?)",
        (fragmento["fuente"], fragmento["titulo"], fragmento["pagina"], fragmento["texto"]),
    )
    db.execute(
        "INSERT INTO vec_fragmentos (rowid, embedding) VALUES (?, ?)",
        (cursor.lastrowid, sqlite_vec.serialize_float32(vector)),
    )


def main():
    parser = argparse.ArgumentParser(
        description="Genera indice.sqlite con la biblioteca de LLMario"
    )
    parser.add_argument(
        "--solo-probar",
        action="store_true",
        help="solo muestra el troceado, sin llamar a la API",
    )
    args = parser.parse_args()

    archivos = list(_archivos())
    if not archivos:
        sys.exit(
            "No hay archivos en biblioteca/ ni docs/ "
            "(formatos: PDF, EPUB, DOCX, DOC, TXT, MD)."
        )

    fragmentos = []
    for ruta in archivos:
        bloques = extraer_archivo(ruta)
        fragmentos.extend(_trocear_bloques(ruta, bloques))

    if args.solo_probar:
        print(f"{len(archivos)} archivo(s), {len(fragmentos)} fragmento(s):")
        for f in fragmentos:
            ubicacion = f"p. {f['pagina']}" if f["pagina"] else f["titulo"]
            print(f"  [{f['fuente']}] {ubicacion} - {len(f['texto'])} caracteres")
        return

    clave = os.environ.get("GEMINI_API_KEY")
    if not clave:
        sys.exit("Falta GEMINI_API_KEY. Exporta la clave y vuelve a intentar.")
    cliente = genai.Client(api_key=clave)

    db = _abrir_db()
    print(f"Embedding de {len(fragmentos)} fragmentos con {MODELO}...")
    for i, fragmento in enumerate(fragmentos, 1):
        vector = embeber(cliente, fragmento["texto"])
        _insertar(db, fragmento, vector)
        if i % 5 == 0 or i == len(fragmentos):
            print(f"  {i}/{len(fragmentos)}")
        time.sleep(0.25)  # cortesía con la cuota gratuita

    db.commit()
    db.close()
    print(
        f"Listo: {ARCHIVO_SQLITE} "
        f"({ARCHIVO_SQLITE.stat().st_size / 1024:.0f} KB, {len(fragmentos)} fragmentos)"
    )


if __name__ == "__main__":
    main()
